"""文档解析器的关键行为测试。"""

from pathlib import Path

import pytest
from docx import Document
from pdfminer.pdfdocument import PDFPasswordIncorrect

from app import parser as parser_module
from app.parser import (
    CorruptedDocumentError,
    DocumentNotFoundError,
    EncryptedPDFError,
    SuspectedScannedPDFError,
    UnsupportedFormatError,
    parse_document,
    smart_truncate,
)


def test_parse_docx_normal(tmp_path: Path) -> None:
    """正常 DOCX 应返回全文和一个逻辑页。"""

    path = tmp_path / "normal.docx"
    document = Document()
    document.add_heading("测试论文", level=1)
    document.add_paragraph("这是一段用于验证 DOCX 解析器的正文。" * 12)
    document.add_table(rows=1, cols=2)
    document.tables[0].cell(0, 0).text = "指标"
    document.tables[0].cell(0, 1).text = "94.2%"
    document.save(path)

    parsed = parse_document(path)

    assert parsed.file_type == "docx"
    assert parsed.page_count == 1
    assert "测试论文" in parsed.text
    assert "指标\t94.2%" in parsed.text


def test_unsupported_format_has_friendly_error(tmp_path: Path) -> None:
    """不支持的格式应抛出专门异常。"""

    path = tmp_path / "notes.txt"
    path.write_text("not supported", encoding="utf-8")

    with pytest.raises(UnsupportedFormatError, match="当前仅支持"):
        parse_document(path)


def test_missing_file_has_friendly_error(tmp_path: Path) -> None:
    """文件不存在时不应落入模糊的通用异常。"""

    with pytest.raises(DocumentNotFoundError, match="文件不存在"):
        parse_document(tmp_path / "missing.pdf")


def test_suspected_scanned_pdf_has_dedicated_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """提取文本少于100字符时应识别为疑似扫描件。"""

    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-fake")

    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakePDF:
        pages = [FakePage()]
        metadata: dict[str, str] = {}

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, *args: object) -> None:
            return None

    monkeypatch.setattr(parser_module.pdfplumber, "open", lambda _: FakePDF())

    with pytest.raises(SuspectedScannedPDFError, match="当前版本不做 OCR"):
        parse_document(path)


def test_encrypted_pdf_has_dedicated_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """加密 PDF 应给出需要密码的明确提示。"""

    path = tmp_path / "encrypted.pdf"
    path.write_bytes(b"%PDF-fake")

    def raise_password_error(_: Path) -> None:
        raise PDFPasswordIncorrect

    monkeypatch.setattr(parser_module.pdfplumber, "open", raise_password_error)

    with pytest.raises(EncryptedPDFError, match="需要密码"):
        parse_document(path)


def test_corrupted_pdf_has_friendly_error(tmp_path: Path) -> None:
    """损坏 PDF 应转换为可理解的项目异常。"""

    path = tmp_path / "broken.pdf"
    path.write_bytes(b"this is not a pdf")

    with pytest.raises(CorruptedDocumentError, match="无法读取"):
        parse_document(path)


def test_smart_truncate_keeps_head_and_tail() -> None:
    """长文本截断后应同时保留开头和结尾。"""

    text = "HEAD-" + ("A" * 200) + ("B" * 200) + "-TAIL"
    result = smart_truncate(text, max_chars=120)

    assert len(result) == 120
    assert result.startswith("HEAD-")
    assert result.endswith("-TAIL")
    assert "中间内容已截断" in result


def test_smart_truncate_returns_short_text_unchanged() -> None:
    """未超长的文本不应被改写。"""

    text = "short paper"

    assert smart_truncate(text, max_chars=100) == text

